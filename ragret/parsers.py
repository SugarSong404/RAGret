"""File parsers for multi-format corpus ingestion."""
from __future__ import annotations

import csv
from abc import ABC, abstractmethod
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Iterable

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from openpyxl import load_workbook


class BaseFileParser(ABC):
    """Parser interface for one or more file suffixes."""

    @property
    @abstractmethod
    def supported_suffixes(self) -> tuple[str, ...]:
        """Return supported suffixes (e.g. (".txt", ".md"))."""

    @abstractmethod
    def parse(self, path: Path) -> list[Document]:
        """Parse one file into one or more langchain Documents."""


class PdfParser(BaseFileParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".pdf",)

    def parse(self, path: Path) -> list[Document]:
        return PyPDFLoader(str(path)).load()


class PlainTextParser(BaseFileParser):
    def __init__(self, suffixes: tuple[str, ...]) -> None:
        self._suffixes = suffixes

    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return self._suffixes

    def parse(self, path: Path) -> list[Document]:
        text = path.read_text(encoding="utf-8")
        return [Document(page_content=text, metadata={"source": str(path)})]


class DocxParser(BaseFileParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".docx",)

    def parse(self, path: Path) -> list[Document]:
        doc = DocxDocument(str(path))
        chunks: list[str] = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                chunks.append(t)
        content = "\n".join(chunks)
        return [Document(page_content=content, metadata={"source": str(path)})]


class XlsxParser(BaseFileParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".xlsx",)

    def parse(self, path: Path) -> list[Document]:
        wb = load_workbook(filename=str(path), data_only=True, read_only=True)
        docs: list[Document] = []
        for sheet in wb.worksheets:
            lines: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(cells):
                    lines.append("\t".join(cells))
            if lines:
                docs.append(
                    Document(
                        page_content="\n".join(lines),
                        metadata={"source": str(path), "sheet": sheet.title},
                    )
                )
        if not docs:
            docs.append(
                Document(
                    page_content="",
                    metadata={"source": str(path)},
                )
            )
        return docs


class CsvParser(BaseFileParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".csv",)

    def parse(self, path: Path) -> list[Document]:
        lines: list[str] = []
        with path.open("r", encoding="utf-8", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                lines.append(", ".join(cell.strip() for cell in row))
        return [Document(page_content="\n".join(lines), metadata={"source": str(path)})]


class HtmlParser(BaseFileParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".html", ".htm")

    def parse(self, path: Path) -> list[Document]:
        raw = path.read_text(encoding="utf-8")
        soup = BeautifulSoup(raw, "html.parser")
        text = soup.get_text(separator="\n", strip=True)
        return [Document(page_content=text, metadata={"source": str(path)})]


class EmlParser(BaseFileParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".eml",)

    def parse(self, path: Path) -> list[Document]:
        with path.open("rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)
        parts: list[str] = []
        subject = str(msg.get("subject", "")).strip()
        if subject:
            parts.append(f"Subject: {subject}")
        sender = str(msg.get("from", "")).strip()
        if sender:
            parts.append(f"From: {sender}")
        to = str(msg.get("to", "")).strip()
        if to:
            parts.append(f"To: {to}")
        body = self._extract_plain_text(msg)
        if body:
            parts.append("")
            parts.append(body)
        return [Document(page_content="\n".join(parts), metadata={"source": str(path)})]

    def _extract_plain_text(self, msg) -> str:
        if msg.is_multipart():
            out: list[str] = []
            for part in msg.walk():
                ctype = part.get_content_type()
                if ctype == "text/plain":
                    txt = part.get_content()
                    if isinstance(txt, str) and txt.strip():
                        out.append(txt.strip())
            return "\n\n".join(out)
        txt = msg.get_content()
        if isinstance(txt, str):
            return txt.strip()
        return ""


def default_parsers() -> list[BaseFileParser]:
    return [
        PdfParser(),
        PlainTextParser((".txt", ".md", ".markdown")),
        DocxParser(),
        XlsxParser(),
        CsvParser(),
        EmlParser(),
        HtmlParser(),
    ]


class ParserRegistry:
    """Simple suffix -> parser registry."""

    def __init__(self, parsers: Iterable[BaseFileParser]) -> None:
        self._by_suffix: dict[str, BaseFileParser] = {}
        for parser in parsers:
            for suffix in parser.supported_suffixes:
                self._by_suffix[suffix.lower()] = parser

    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return tuple(sorted(self._by_suffix.keys()))

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self._by_suffix

    def parse_file(self, path: Path) -> list[Document]:
        parser = self._by_suffix.get(path.suffix.lower())
        if parser is None:
            raise ValueError(
                f"Unsupported file type: {path.suffix}. "
                f"Supported: {', '.join(self.supported_suffixes)}"
            )
        return parser.parse(path)
