"""Document parser abstractions and concrete parser implementations."""
from __future__ import annotations

import csv
from abc import ABC, abstractmethod
from email import policy
from email.parser import BytesParser
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxFile
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from openpyxl import load_workbook


class DocumentParser(ABC):
    """Base class for all file parsers used by indexing."""

    @property
    @abstractmethod
    def supported_suffixes(self) -> tuple[str, ...]:
        """Lower-cased file suffixes this parser can read."""

    @abstractmethod
    def parse(self, path: Path) -> list[Document]:
        """Parse one file into langchain documents."""


class PdfParser(DocumentParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".pdf",)

    def parse(self, path: Path) -> list[Document]:
        return PyPDFLoader(str(path)).load()


class TextLikeParser(DocumentParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".txt", ".md", ".markdown")

    def parse(self, path: Path) -> list[Document]:
        return TextLoader(str(path), encoding="utf-8").load()


class CsvParser(DocumentParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".csv",)

    def parse(self, path: Path) -> list[Document]:
        rows: list[str] = []
        with path.open("r", encoding="utf-8", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append("\t".join(str(cell or "") for cell in row))
        text = "\n".join(rows).strip()
        return [Document(page_content=text, metadata={"source": str(path)})]


class HtmlParser(DocumentParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".html", ".htm")

    def parse(self, path: Path) -> list[Document]:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(raw, "html.parser")
        text = soup.get_text(separator="\n", strip=True)
        return [Document(page_content=text, metadata={"source": str(path)})]


class EmlParser(DocumentParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".eml",)

    def parse(self, path: Path) -> list[Document]:
        with path.open("rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)
        chunks: list[str] = []
        subject = str(msg.get("subject") or "").strip()
        if subject:
            chunks.append(f"Subject: {subject}")
        for part in msg.walk():
            ctype = part.get_content_type()
            if ctype == "text/plain":
                body = part.get_content()
                if isinstance(body, str) and body.strip():
                    chunks.append(body.strip())
            elif ctype == "text/html":
                html = part.get_content()
                if isinstance(html, str) and html.strip():
                    text = BeautifulSoup(html, "html.parser").get_text(separator="\n", strip=True)
                    if text:
                        chunks.append(text)
        text = "\n\n".join(chunks).strip()
        return [Document(page_content=text, metadata={"source": str(path)})]


class DocxParser(DocumentParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".docx",)

    def parse(self, path: Path) -> list[Document]:
        doc = DocxFile(str(path))
        paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        text = "\n".join(paras).strip()
        return [Document(page_content=text, metadata={"source": str(path)})]


class XlsxParser(DocumentParser):
    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return (".xlsx",)

    def parse(self, path: Path) -> list[Document]:
        wb = load_workbook(str(path), data_only=True, read_only=True)
        sections: list[str] = []
        for ws in wb.worksheets:
            lines: list[str] = []
            for row in ws.iter_rows(values_only=True):
                vals = [str(v).strip() for v in row if v is not None and str(v).strip() != ""]
                if vals:
                    lines.append("\t".join(vals))
            if lines:
                sections.append(f"[Sheet: {ws.title}]\n" + "\n".join(lines))
        text = "\n\n".join(sections).strip()
        return [Document(page_content=text, metadata={"source": str(path)})]


PARSERS: tuple[DocumentParser, ...] = (
    PdfParser(),
    TextLikeParser(),
    CsvParser(),
    HtmlParser(),
    EmlParser(),
    DocxParser(),
    XlsxParser(),
)

SUPPORTED_SUFFIXES: tuple[str, ...] = tuple(
    sorted({suf for parser in PARSERS for suf in parser.supported_suffixes})
)


def parse_document(path: Path) -> list[Document]:
    suf = path.suffix.lower()
    for parser in PARSERS:
        if suf in parser.supported_suffixes:
            return parser.parse(path)
    allowed = ", ".join(SUPPORTED_SUFFIXES)
    raise ValueError(f"Unsupported file type: {path.suffix}. Supported: {allowed}")
